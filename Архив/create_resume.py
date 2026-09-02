#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генератор резюме Миронова А.С. — CPO"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, color=(0, 51, 102)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(*color)
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_experience_block(doc, period, company, position, responsibilities, results):
    """Добавляет блок опыта: сначала результат(ы), потом обязанности"""
    # Период и должность
    p = doc.add_paragraph()
    run_period = p.add_run(period)
    run_period.font.size = Pt(10)
    run_period.font.name = "Calibri"
    run_period.font.color.rgb = RGBColor(100, 100, 100)
    run_period.italic = True

    p_company = doc.add_paragraph()
    run_company = p_company.add_run(company)
    run_company.bold = True
    run_company.font.size = Pt(12)
    run_company.font.name = "Calibri"
    run_company.font.color.rgb = RGBColor(0, 0, 0)

    run_position = p_company.add_run(" — " + position)
    run_position.font.size = Pt(11)
    run_position.font.name = "Calibri"

    # Результаты (сначала!)
    if results:
        p_results = doc.add_paragraph()
        run_results_label = p_results.add_run("Достижения: ")
        run_results_label.bold = True
        run_results_label.font.size = Pt(11)
        run_results_label.font.name = "Calibri"
        
        for result in results:
            p_result = doc.add_paragraph(style="List Bullet")
            run_result = p_result.add_run(result)
            run_result.font.size = Pt(11)
            run_result.font.name = "Calibri"

    # Обязанности (потом)
    if responsibilities:
        p_resp_label = doc.add_paragraph()
        run_resp_label = p_resp_label.add_run("Обязанности: ")
        run_resp_label.bold = True
        run_resp_label.font.size = Pt(11)
        run_resp_label.font.name = "Calibri"
        
        for resp in responsibilities:
            p_resp = doc.add_paragraph(style="List Bullet")
            run_resp = p_resp.add_run(resp)
            run_resp.font.size = Pt(11)
            run_resp.font.name = "Calibri"

def create_resume():
    doc = Document()
    
    # Настройки документа
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ===== ЗАГОЛОВОК =====
    p_name = doc.add_paragraph()
    run_name = p_name.add_run("Миронов Андрей Сергеевич")
    run_name.bold = True
    run_name.font.size = Pt(24)
    run_name.font.name = "Calibri"
    run_name.font.color.rgb = RGBColor(0, 51, 102)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_age = doc.add_paragraph()
    run_age = p_age.add_run("42 года, родился 1 апреля 1984")
    run_age.font.size = Pt(11)
    run_age.font.name = "Calibri"
    run_age.font.color.rgb = RGBColor(80, 80, 80)
    p_age.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Контакты
    p_contact = doc.add_paragraph()
    run_contact = p_contact.add_run("+7 (916) 1586826  |  sonic.kineshma@gmail.com")
    run_contact.font.size = Pt(11)
    run_contact.font.name = "Calibri"
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # пустая строка

    # ===== ЖЕЛАЕМАЯ ДОЛЖНОСТЬ =====
    p_target = doc.add_paragraph()
    run_target_label = p_target.add_run("Желаемая должность: ")
    run_target_label.bold = True
    run_target_label.font.size = Pt(14)
    run_target_label.font.name = "Calibri"
    
    run_target = p_target.add_run("CPO (Chief Product Officer)")
    run_target.font.size = Pt(14)
    run_target.font.name = "Calibri"
    run_target.font.color.rgb = RGBColor(0, 51, 102)

    # Специфика
    p_spec = doc.add_paragraph()
    run_spec = p_spec.add_run("Специализация: Менеджмент продуктов, AI/ML-продукты, B2B-продукты")
    run_spec.font.size = Pt(11)
    run_spec.font.name = "Calibri"

    p_loc = doc.add_paragraph()
    run_loc = p_loc.add_run("Готов к командировкам, релокация")
    run_loc.font.size = Pt(11)
    run_loc.font.name = "Calibri"

    doc.add_paragraph()

    # ===== ОПЫТ РАБОТЫ =====
    add_section_title(doc, "ОПЫТ РАБОТЫ — 13 лет 8 месяцев")
    
    # --- ДОМ.РФ ---
    add_experience_block(
        doc,
        "Февраль 2023 — Февраль 2026 (3 года 1 месяц)",
        "Банк ДОМ.РФ", "domrfbank.ru",
        responsibilities=[
            "Управление продуктовой командой (бизнес-аналитики, разработчики, тестировщики, метрологи)",
            "Руководство数字化-проектами (системы биометрической идентификации, компьютерного зрения, метрологии, биометрические платформы)",
            "Запуск AI-продуктов по направлениям ИИТ (CV, ML, GenAI)",
            "Управление продуктовой разработкой (включая собственную LLM)",
            "Организация R&D-студии (проектирование новых направлений, архитектура, R&D-процессы)",
            "Запуск B2B-направления: кредитный скоринг и финтех-решения для корпоративных клиентов",
            "Продуктовый запуск ИИТ-платформы (найм ML-инженеров, организация работы с финтех-партнёрами)",
            "Управление бюджетом проектов (в том числе собственные ML-модели финтех-продуктов)",
        ],
        results=[
            "Собрал и выстроил работу продуктовой команды из 380+ человек в направлении AI/ML (CV, ML, GenAI)",
            "Организовал R&D-процесс для создания собственной LLM (финтех-модель)",
            "Запустил B2B-решение (кредитный скоринг) — сокращение времени скоринга с 3 месяцев до 1-2 недель",
            "Обеспечил финансовый эффект: снижение операционных расходов на 30 млн руб. в год (кредитный скоринг), экономия 10 млн руб. в год (автоматизация)",
            "Организовал Agile-процессы для 250 человек из 8 месяцев",
            "Выстроил R&D-студию: запустил 6 MVP за 3 месяца (система верификации, скоринг, финтех-модели)",
        ]
    )

    # --- Сбер: Ведущий продукт ---
    add_experience_block(
        doc,
        "Август 2018 — Март 2021 (2 года 8 месяцев)",
        "Сбер", "rabota.sber.ru",
        responsibilities=[
            "Управление продуктом 'Примитивные технологии искусственного интеллекта'",
            "Управление коммерческими проектами (анализ рынка, формирование roadmap)",
            "Анализ конкурентов и партнёров",
            "Управление продуктовой командой (аналитики, разработчики, ML-инженеры)",
            "Управление техническим заданием и приоритизация задач",
            "Управление процессом разработки: от сбора требований до запуска",
        ],
        results=[
            "Вырос из Product Owner в Ведущего продукта за 5 проектов",
            "Увеличил выручку на 15% за счёт новых телематических предложений",
            "Запустил MVP продукта (телематика, эквайринг, реклама)",
            "Выстроил Agile-процессы для 250 человек за 8 месяцев",
            "Увеличил покрытие телематикой с 1 млн до 3 млн пользователей за 3 месяца",
            "Достиг роста выручки с 30 до 50 млн руб. в год по B2B-направлению",
        ]
    )

    # --- Сбер: Product Owner ---
    add_experience_block(
        doc,
        "Август 2018 — Март 2021",
        "Сбер", "Product Owner проекта: 'Примитивные технологии искусственного интеллекта'",
        responsibilities=[
            "Управление требованиями и приоритизация задач",
            "Обратная связь от клиентов: от сбора до реализации",
            "Управление продуктовой командой (аналитики, разработчики)",
            "Взаимодействие со стейкхолдерами",
        ],
        results=[
            "Запустил MVP продукта за 3 месяца",
            "Обеспечил рост выручки B2B-направления с 30 до 50 млн руб./год",
            "Вырос до позиции Ведущего продукта",
        ]
    )

    # --- Сбер: Старший системный аналитик ---
    add_experience_block(
        doc,
        "Март 2017 — Март 2021 (4 года 1 месяц)",
        "Сбер", "Старший системный аналитик",
        responsibilities=[
            "Анализ бизнес-требований",
            "Написание ТЗ и спецификаций",
            "Проектирование интеграций",
            "Управление командой аналитиков",
            "Взаимодействие с разработчиками и QA",
            "Управление процессом разработки: от сбора требований до запуска",
        ],
        results=[
            "Собрал команду аналитиков из 5 человек",
            "Увеличил покрытие IoT-устройствами с 1 млн до 3 млн за 3 месяца",
            "Обеспечил рост выручки с 30 до 50 млн руб./год",
            "Вырос до позиции Product Owner",
        ]
    )

    # --- ЮЭйКьюБ ---
    add_experience_block(
        doc,
        "Август 2016 — Март 2017 (8 месяцев)",
        "ООО ЮЭйКьюБ", "www.ue-cb.com",
        responsibilities=[
            "Проектирование и внедрение решений",
            "Управление проектами (ТЗ, спецификации, интеграции)",
            "Управление командой разработки",
            "Взаимодействие с заказчиками",
            "Обеспечение качества решений",
        ],
        results=[
            "Завершил проект для ФНС (электронная отчётность) в срок",
            "Успешно сдал проект по внедрению решений для ФНС",
        ]
    )

    doc.add_paragraph()

    # ===== ОБРАЗОВАНИЕ =====
    add_section_title(doc, "ОБРАЗОВАНИЕ")
    p_edu = doc.add_paragraph()
    run_edu = p_edu.add_run("Московский государственный университет экономики, статистики и информатики (МЭСИ)")
    run_edu.bold = True
    run_edu.font.size = Pt(11)
    run_edu.font.name = "Calibri"
    
    p_edu_year = doc.add_paragraph()
    run_edu_year = p_edu_year.add_run("2007 — Высшее")
    run_edu_year.font.size = Pt(11)
    run_edu_year.font.name = "Calibri"

    doc.add_paragraph()

    # ===== НАВЫКИ =====
    add_section_title(doc, "КЛЮЧЕВЫЕ НАВЫКИ")
    
    skills = [
        "Agile/Project Management",
        "Бизнес-анализ",
        "Бюджетирование",
        "Бизнес-моделирование",
        "Scrum",
        "Управление продуктом — полный цикл",
        "Управление требованиями",
        "Управление командой",
        "Постановка задач команде",
        "Управление коммерческими проектами",
        "PMBOK",
        "Организация наставничества",
        "Unit-экономика",
        "Прототипирование",
    ]
    
    for skill in skills:
        p_skill = doc.add_paragraph(style="List Bullet")
        run_skill = p_skill.add_run(skill)
        run_skill.font.size = Pt(11)
        run_skill.font.name = "Calibri"

    doc.add_paragraph()

    # ===== ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ =====
    add_section_title(doc, "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ")
    
    p_lang = doc.add_paragraph()
    run_lang = p_lang.add_run("Языки: ")
    run_lang.bold = True
    run_lang.font.size = Pt(11)
    run_lang.font.name = "Calibri"
    run_lang2 = p_lang.add_run("Русский — Родной")
    run_lang2.font.size = Pt(11)
    run_lang2.font.name = "Calibri"

    doc.add_paragraph()

    # ===== О СЕБЕ =====
    add_section_title(doc, "О СЕБЕ")
    
    about_text = """Product Owner и CPO с 10-летним стажем в крупных финтех-компаниях (Сбер, Банк ДОМ.РФ).

Запускал AI/ML-продукты (CV, ML, GenAI) — от идеи до production. Умею выстраивать продуктовые команды с нуля (собирал команду из 380+ человек), выстраивать R&D-процессы и запускать B2B-решения.

Основные компетенции:
• Управление продуктовой разработкой (full cycle): от сбора требований до запуска и масштабирования
• AI/ML-продукты: организация R&D, работа с данными, внедрение ML-моделей
• Управление командой: найм, мотивация, развитие, наставничество
• Бюджетирование и управление коммерческими проектами (Unit-экономика)
• Agile/Scrum: выстраивал процессы для 250+ человек за 8 месяцев"""

    p_about = doc.add_paragraph()
    run_about = p_about.add_run(about_text)
    run_about.font.size = Pt(11)
    run_about.font.name = "Calibri"

    # Сохранение
    output_path = "/home/user/Миронов_Андрей_CPO_Резюме.docx"
    doc.save(output_path)
    print(f"Резюме сохранено: {output_path}")
    return output_path

if __name__ == "__main__":
    create_resume()
