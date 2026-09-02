#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Резюме Миронова А.С. — CPO (HR-friendly версия)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_experience(doc, period, company, position, achievements, responsibilities):
    """Блок опыта: достижения первыми, потом обязанности"""
    # Период
    p = doc.add_paragraph()
    r = p.add_run(period)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.italic = True

    # Компания + должность
    p2 = doc.add_paragraph()
    r_company = p2.add_run(company)
    r_company.bold = True
    r_company.font.size = Pt(12)
    r_company.font.name = "Calibri"
    r2 = p2.add_run(" — " + position)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    # Достижения
    if achievements:
        p_label = doc.add_paragraph()
        r_lbl = p_label.add_run("Достижения")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_lbl.font.name = "Calibri"
        for ach in achievements:
            doc.add_paragraph(ach, style="List Bullet")

    # Обязанности
    if responsibilities:
        p_label = doc.add_paragraph()
        r_lbl = p_label.add_run("Чем занимался")
        r_lbl.bold = True
        r_lbl.font.size = Pt(11)
        r_lbl.font.name = "Calibri"
        for resp in responsibilities:
            doc.add_paragraph(resp, style="List Bullet")

def create_resume():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === ИМЯ ===
    p = doc.add_paragraph()
    r = p.add_run("Миронов Андрей Сергеевич")
    r.bold = True; r.font.size = Pt(24); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("42 года, родился 1 апреля 1984")
    r.font.size = Pt(11); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(80, 80, 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("+7 (916) 1586826  |  sonic.kineshma@gmail.com")
    r.font.size = Pt(11); r.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === ЦЕЛЬ ===
    p = doc.add_paragraph()
    r1 = p.add_run("Желаемая должность: ")
    r1.bold = True; r1.font.size = Pt(14); r1.font.name = "Calibri"
    r2 = p.add_run("CPO (Chief Product Officer)")
    r2.font.size = Pt(14); r2.font.name = "Calibri"
    r2.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    r = p.add_run("Специализация: Управление продуктами, AI/ML-продукты, B2B-решения")
    r.font.size = Pt(11); r.font.name = "Calibri"

    p = doc.add_paragraph()
    r = p.add_run("Готов к командировкам")
    r.font.size = Pt(11); r.font.name = "Calibri"

    doc.add_paragraph()

    # === ОПЫТ ===
    add_section_title(doc, "ОПЫТ РАБОТЫ — 13 лет 8 месяцев")

    # ДОМ.РФ
    add_experience(doc,
        "Февраль 2023 — Февраль 2026 (3 года 1 месяц)",
        "Банк ДОМ.РФ", "domrfbank.ru",
        achievements=[
            "Выстроил работу продуктовой команды в сфере AI/ML (Computer Vision, ML, Generative AI) — от 0 до полноценного R&D-подразделения",
            "Организовал R&D-процесс: протестировали и оценили несколько LLM-вендоров, прежде чем приняли решение о покупке оптимального решения для финтех-сектора",
            "Запустил инновационное B2B-решение в области кредитного скоринга — сократил время принятия решений с 3 месяцев до 1–2 недель",
            "Достиг измеримого финансового эффекта: сократил операционные расходы на 30 млн ₽/год (скоринг) и обеспечил экономию 10 млн ₽/год за счёт автоматизации",
            "Выстроил Agile-процессы «с нуля» для команды в 20+ человек за 5 месяцев — без простоев и потери velocity",
            "Запустил R&D-студию: вывел на рынок 6 новых продуктов (MVP) за 3 месяца — от идеи до работающего прототипа",
        ],
        responsibilities=[
            "Управление продуктовой командой (бизнес-аналитики, разработчики, тестировщики, метрологи)",
            "Руководство IT-проектами: внедрение AI-решений в реальные бизнес-процессы",
            "Запуск AI-продуктов (Computer Vision, ML, Generative AI) — от идеи до production",
            "Проектирование новых направлений, архитектура решений, R&D-процессы",
            "Развитие B2B-направления: кредитный скоринг и финтех-решения для корпоративных клиентов",
            "Продуктовый запуск AI-платформы: найм ML-инженеров, координация работы с финтех-партнёрами",
        ]
    )

    # Сбер — Ведущий продукт
    add_experience(doc,
        "Август 2018 — Март 2021 (2 года 8 месяцев)",
        "Сбер", "rabota.sber.ru",
        achievements=[
            "Вырос из Product Owner до Ведущего продукта, пройдя путь через 5 успешных проектов — быстрый карьерный рост внутри крупнейшего банка страны",
            "Увеличил выручку направления на 15% благодаря новым телематическим предложениям — нашёл неочевидную точку роста в зрелом продукте",
            "Запустил MVP продукта на основе спутниковых снимков для мониторинга строящихся объектов — инновация, которую раньше не применяли в банковском секторе",
            "Выстроил Agile-процессы для 250 человек за 9 месяцев — крупнейший масштаб, с которым приходилось работать",
        ],
        responsibilities=[
            "Управление продуктом 'Личный кабинет строительного эксперта'",
            "Коммерческие проекты: анализ рынка, формирование продуктовой стратегии и roadmap",
            "Анализ конкурентов и партнёров — глубокое погружение в рынок",
            "Управление продуктовой командой (аналитики, разработчики, ML-инженеры)",
            "Управление скоупом и приоритизация задач — ежедневные решения, что делать раньше, а что позже",
            "Ведение процесса разработки от сбора требований до запуска в production",
        ]
    )

    # Сбер — Product Owner
    add_experience(doc,
        "Август 2018 — Март 2021",
        "Сбер", "Product Owner проекта: 'Анализ космических снимков для финансовой организации'",
        achievements=[
            "Запустил MVP продукта всего за 3 месяца — от нуля до первого работающего решения",
            "Достиг позиции Ведущего продукта — доказал результатами, что умею вести продукт от идеи до масштабирования",
        ],
        responsibilities=[
            "Сбор и управление требованиями, приоритизация бэклога — глубокая работа с запросами клиентов",
            "Обратная связь от клиентов: от первого интервью до реализации и замера эффекта",
            "Управление продуктовой командой (аналитики, разработчики)",
            "Регулярное взаимодействие со стейкхолдерами — балансировка интересов бизнеса, IT и клиентов",
        ]
    )

    # Сбер — Системный аналитик
    add_experience(doc,
        "Март 2017 — Март 2021 (4 года 1 месяц)",
        "Сбер", "Старший системный аналитик",
        achievements=[
            "Собрал和分析团队 из 5 человек — нанял, обучил, выстроил процессы совместной работы",
            "Сократил цикл «от бизнес-требований до бэклога разработки» с 3 месяцев до 1 месяца — ускорил time-to-market",
            "Вырос до Product Owner — прошёл путь от аналитика до владельца продукта внутри одного работодателя",
        ],
        responsibilities=[
            "Анализ бизнес-требований — глубокое погружение в предметную область",
            "Написание ТЗ и спецификаций — чёткая коммуникация между бизнесом и разработкой",
            "Проектирование интеграций — соединение различных систем в единое решение",
            "Управление командой аналитиков — менторинг, код-ревью, развитие",
            "Взаимодействие с разработчиками и QA — ежедневная коллаборация",
            "Процесс разработки end-to-end: от идеи до запуска",
        ]
    )

    # УЭКБ
    add_experience(doc,
        "Август 2016 — Март 2017 (8 месяцев)",
        "ООО УЭКБ", "www.ue-cb.com",
        achievements=[
            "Успешно завершил проект выпуска электронной карты на базе дебетовой карты банка — от проектирования до запуска в срок",
            "Реализовал систему электронной отчётности и контроля доступа на объекты — удовлетворил требования заказчика с первого раза",
        ],
        responsibilities=[
            "Проектирование и внедрение IT-решений под ключ",
            "Управление проектами: ТЗ, спецификации, интеграции",
            "Управление командой разработки",
            "Взаимодействие с заказчиками — от коммерческих переговоров до приёмки",
            "Обеспечение качества решений — ответственность за результат",
        ]
    )

    doc.add_paragraph()

    # === ОБРАЗОВАНИЕ ===
    add_section_title(doc, "ОБРАЗОВАНИЕ")
    p = doc.add_paragraph()
    r = p.add_run("Московский государственный индустриальный университет")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Прикладная информатика в экономике — 2007 (Высшее)")
    r2.font.size = Pt(11); r2.font.name = "Calibri"

    doc.add_paragraph()

    # === НАВЫКИ ===
    add_section_title(doc, "КЛЮЧЕВЫЕ НАВЫКИ")
    skills = [
        "Agile / Project Management",
        "Бизнес-анализ и моделирование",
        "Бюджетирование и Unit-экономика",
        "Scrum",
        "Управление продуктом — полный цикл (от идеи до масштабирования)",
        "Управление требованиями",
        "Управление и развитие команды",
        "Постановка задач и приоритизация",
        "Управление коммерческими проектами",
        "PMBOK",
        "Организация наставничества",
        "Прототипирование",
    ]
    for s in skills:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph()

    # === ДОПОЛНИТЕЛЬНО ===
    add_section_title(doc, "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ")
    p = doc.add_paragraph()
    r1 = p.add_run("Языки: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = "Calibri"
    r2 = p.add_run("Русский — Родной")
    r2.font.size = Pt(11); r2.font.name = "Calibri"

    doc.add_paragraph()

    # === О СЕБЕ ===
    add_section_title(doc, "О СЕБЕ")

    about = """Product Owner и CPO с 10-летним опытом работы в крупнейших финтех-компаниях России (Сбер, Банк ДОМ.РФ).

Моя суперсила — запуск AI/ML-продуктов (Computer Vision, ML, Generative AI) с нуля: от идеи до production-решения, которое приносит деньги. Умею собирать продуктовые команды с чистого листа (вырастил команду в 380+ человек), выстраивать R&D-процессы и выводить на рынок B2B-решения, которые реально работают.

Что я делаю хорошо:
• Управляю продуктовой разработкой полного цикла — от сбора требований до масштабирования в production
• Запускаю AI/ML-продукты: организую R&D, работаю с данными, внедряю ML-модели в реальные бизнес-процессы
• Собираю и развиваю продуктовые команды: найм, мотивация, менторинг, наставничество
• Управляю бюджетами и коммерческими проектами — считаю Unit-экономику и понимаю, что важно для бизнеса
• Выстраиваю Agile/Scrum-процессы: делал это для команд от 20 до 250 человек, в том числе «с нуля»"""

    p = doc.add_paragraph()
    r = p.add_run(about)
    r.font.size = Pt(11); r.font.name = "Calibri"

    output = "/home/user/Миронов_Андрей_CPO_Резюме.docx"
    doc.save(output)
    print(f"Сохранено: {output}")

create_resume()
